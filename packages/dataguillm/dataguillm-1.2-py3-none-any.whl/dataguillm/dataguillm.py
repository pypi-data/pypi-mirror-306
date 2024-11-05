import pandas as pd
import seaborn as sns
import plotly.express as px
import numpy as np
import warnings
warnings.filterwarnings('ignore')
from scipy.stats import chi2_contingency
from scipy import stats
import statsmodels.api as sm
from statsmodels.formula.api import ols
import datapane as dp
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.preprocessing import LabelEncoder
from sklearn.model_selection import train_test_split
from groq import Groq
from docx import Document
from docx.shared import Inches

class AnalyzeData:
    def __init__(self, df, target_col, api_key):
        self.api_key= api_key
        self.client = Groq(
                        api_key= self.api_key,
                    )
        self.df = df.copy()
        self.target_col = target_col
        # Determine if the target column is categorical or continuous
        if pd.api.types.is_numeric_dtype(self.df[self.target_col]):
            if self.df[self.target_col].nunique() < 10:
                self.df[self.target_col] = self.df[self.target_col].astype('category')
                self.target_type = 'categorical'
            else:
                self.target_type = 'continuous'
        else:
            self.target_type = 'categorical'

        if self.df.shape[0] > 2500000:
            if self.target_type == "categorical" or self.target_type == "category":
                self.df, _= train_test_split(self.df, train_size=2500000,
                                                stratify=self.df[self.target_col],
                                                random_state=42,)
            else:
                self.df= self.df.head(2500000)
        # Identify datetime columns
        self.datetime_cols = self.df.select_dtypes(include=['datetime64[ns]', 'datetime64[ns, UTC]']).columns.tolist()
        # Attempt to parse object columns that might be datetime
        object_cols = self.df.select_dtypes(include=['object']).columns.tolist()
        for col in object_cols:
            if col not in self.datetime_cols and col != self.target_col:
                try:
                    parsed_col = pd.to_datetime(self.df[col], errors='raise')
                    self.df[col] = parsed_col
                    self.datetime_cols.append(col)
                except (ValueError, TypeError):
                    continue
        # Note: We no longer drop datetime columns

    def chat_llm(self, data):
        try:
            chat_completion = self.client.chat.completions.create(
                                messages=[
                                    {
                                        "role": "user",
                                        "content": f"Explain this {data} to me for a short report",
                                    }
                                ],
                                model="llama-3.2-90b-vision-preview",
                            )
            result= chat_completion.choices[0].message.content
        except:
            result= "Could not generate content from LLM. Please try again!"
        return result

    def __box_plot_outliers(self):
        figures = []
        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        if self.target_col in numeric_cols:
            numeric_cols.remove(self.target_col)
        for col in numeric_cols:
            fig = px.box(
                self.df,
                y=col,
                title=f'Box plot of {col}',
                template='plotly_white',
                color_discrete_sequence=px.colors.sequential.Greys_r
            )
            fig.update_layout(title_x=0.5)  # Center-align the title
            figures.append(fig)
        return figures

    def __class_imbalance(self):
        if self.target_type == 'categorical':
            class_counts = self.df[self.target_col].value_counts().to_frame().reset_index().rename(columns={"index": self.target_col, self.target_col: "imbalance count"})
            return class_counts
        else:
            class_counts = pd.DataFrame({"column": "No value count of target column as it is not categorical"}, index=[0])
            return class_counts

    def __decision_tree_feature_importance(self):
        df= self.df.copy()
        df= df.dropna()
        X = df.drop(columns=[self.target_col])
        y = df[self.target_col]
        # Exclude datetime columns from X
        X = X.drop(columns=self.datetime_cols)
        # Handle categorical variables in features
        le = LabelEncoder()
        X_encoded = X.copy()
        categorical_cols = X_encoded.select_dtypes(include=['object', 'category']).columns.tolist()
        for col in categorical_cols:
            X_encoded[col] = le.fit_transform(X_encoded[col].astype(str))
        # No need to handle datetime columns, as they are excluded

        if self.target_type == 'categorical':
            # Encode categorical target variable
            y_encoded = le.fit_transform(y)
            model = RandomForestClassifier(random_state=42, n_estimators=100)
        else:
            y_encoded = y
            model = RandomForestRegressor(random_state=42, n_estimators=100)
        # Fit the model
        model.fit(X_encoded, y_encoded)
        # Get feature importances
        importances = model.feature_importances_
        feature_names = X_encoded.columns
        feature_importance_df = pd.DataFrame({
            'Feature': feature_names,
            'Importance': importances
        }).sort_values(by='Importance', ascending=False)
        return feature_importance_df

    def __decision_tree_gradientboosting_feature_importance(self):
        df= self.df.copy()
        df= df.dropna()
        X = df.drop(columns=[self.target_col])
        y = df[self.target_col]
        # Exclude datetime columns from X
        X = X.drop(columns=self.datetime_cols)
        # Handle categorical variables in features
        le = LabelEncoder()
        X_encoded = X.copy()
        categorical_cols = X_encoded.select_dtypes(include=['object', 'category']).columns.tolist()
        for col in categorical_cols:
            X_encoded[col] = le.fit_transform(X_encoded[col].astype(str))
        # No need to handle datetime columns, as they are excluded

        if self.target_type == 'categorical':
            # Encode categorical target variable
            y_encoded = le.fit_transform(y)
            model = DecisionTreeClassifier(random_state=42, max_depth=2)
        else:
            y_encoded = y
            model = DecisionTreeRegressor(random_state=42, max_depth=2)
        # Fit the model
        model.fit(X_encoded, y_encoded)
        # Get feature importances
        importances = model.feature_importances_
        feature_names = X_encoded.columns
        feature_importance_df = pd.DataFrame({
            'Feature': feature_names,
            'Importance': importances
        }).sort_values(by='Importance', ascending=False)
        return feature_importance_df

    def univariate_analysis(self):
        shape_rows= self.df.shape[0]
        shape_columns= self.df.shape[1]
        if len(self.datetime_cols) > 0:
            shape_df = f"{shape_rows} rows and {shape_columns} columns" + " and the datetime column(s) is/are: " + str(self.datetime_cols)
        else:
            shape_df = f"{shape_rows} rows and {shape_columns} columns"
        unique_df_count = self.df.nunique().to_frame().reset_index().rename(columns={"index": "column", 0: "unique count"})
        describe_df = self.df.describe(include='all').reset_index()
        dtypes_df = self.df.dtypes.to_frame().T
        na_df = self.df.isna().sum().to_frame().T

        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_columns = self.df.select_dtypes(include=['object', 'category', 'bool']).columns.tolist()

        # Remove target column from numeric_cols and categorical_columns
        if self.target_col in numeric_cols:
            numeric_cols.remove(self.target_col)
        if self.target_col in categorical_columns:
            categorical_columns.remove(self.target_col)

        # Remove datetime columns from categorical_columns
        for dt_col in self.datetime_cols:
            if dt_col in categorical_columns:
                categorical_columns.remove(dt_col)

        df_cat = None
        if len(categorical_columns) > 0:
            df_cat = self.df[categorical_columns].apply(lambda x: x.value_counts()).T.stack().reset_index().rename(columns={'level_0': 'column', 'level_1': 'value', 0: 'count'})

        # Outlier analysis using IQR method
        outlier_info = []
        if len(numeric_cols) > 0:
            for col in numeric_cols:
                Q1 = self.df[col].quantile(0.25)
                Q3 = self.df[col].quantile(0.75)
                IQR = Q3 - Q1
                lower_bound = Q1 - 1.5 * IQR
                upper_bound = Q3 + 1.5 * IQR
                outliers = self.df[(self.df[col] < lower_bound) | (self.df[col] > upper_bound)]
                outlier_info.append({
                    'column': col,
                    'number_of_outliers': outliers.shape[0],
                    'lower_bound': lower_bound,
                    'upper_bound': upper_bound
                })
            outlier_df = pd.DataFrame(outlier_info)
        else:
            outlier_df = None

        figures = []
        box_figures = []

        # Plotting numeric columns
        if len(numeric_cols) > 0:
            for col in numeric_cols:
                fig = px.histogram(
                    self.df,
                    x=col,
                    nbins=30,
                    title=f'Histogram of {col}',
                    labels={col: col},
                    template='plotly_white',
                )
                fig.update_layout(title_x=0.5)  # Center-align the title
                figures.append(fig)
            box_figures = self.__box_plot_outliers()

        # Plotting target column
        target_figures = []
        if self.target_type == 'categorical':
            fig = px.bar(
                self.df[self.target_col].value_counts().reset_index(),
                x='index',
                y=self.target_col,
                color='index',
                title=f'Class imbalance plot of {self.target_col}',
                template='plotly_white',
                color_discrete_sequence=px.colors.sequential.Blues
            )
            fig.update_layout(title_x=0.5)  # Center-align the title
            target_figures.append(fig)
        else:
            fig = px.histogram(
                self.df,
                x=self.target_col,
                nbins=30,
                title=f'Histogram of {self.target_col}',
                labels={self.target_col: self.target_col},
                template='plotly_white',
            )
            fig.update_layout(title_x=0.5)  # Center-align the title
            target_figures.append(fig)
            fig_box = px.box(
                self.df,
                y=self.target_col,
                title=f'Box plot of {self.target_col}',
                template='plotly_white',
                color_discrete_sequence=px.colors.sequential.Greys_r
            )
            fig_box.update_layout(title_x=0.5)  # Center-align the title
            target_figures.append(fig_box)

        # Class imbalance
        pie_figures = self.__class_imbalance()

        return (
            shape_df,
            describe_df,
            dtypes_df,
            na_df,
            df_cat,
            outlier_df,
            target_figures,
            figures,
            box_figures,
            pie_figures,
            unique_df_count
        )

    def multivariate_analysis(self):
        target_column = self.target_col
        df_category = self.df.copy()
        chi_square_results = []
        anova_results = []
        plots = []
        corr_plots = []
        time_series_figures = []  # Moved time_series_figures here
        response_llm_correlation= []
        corr_tables= []

        # Exclude datetime columns
        datetime_cols = self.datetime_cols

        if self.target_type == 'categorical':
            # Handle the case where target is categorical
            categorical_columns = list(df_category.select_dtypes(include=['object', 'category', 'bool']))
            if target_column in categorical_columns:
                categorical_columns.remove(target_column)
            # Exclude datetime columns
            categorical_columns = [col for col in categorical_columns if col not in datetime_cols]

            numerical_cols = list(df_category.select_dtypes(include=[np.number]))
            # Exclude datetime columns
            numerical_cols = [col for col in numerical_cols if col not in datetime_cols]

            # Chi-Square Tests
            for col in categorical_columns:
                crosstab = pd.crosstab(df_category[target_column], df_category[col])
                if crosstab.shape[0] < 2 or crosstab.shape[1] < 2:
                    continue  # Skip if not enough categories for chi-square test
                chi2, p, _, _ = chi2_contingency(crosstab)
                significance = 'Significant' if p < 0.05 else 'Not Significant'
                chi_square_results.append({
                    'Column': col,
                    'Association': significance,
                    'p-value': round(p, 3)
                })

            # ANOVA Tests
            for intcol in numerical_cols:
                if df_category[target_column].nunique() < 2:
                    continue  # Skip if target has less than 2 categories
                model = ols(f'{intcol} ~ C({target_column})', data=df_category).fit()
                anova_table = sm.stats.anova_lm(model, typ=1)
                anova_pvalue = anova_table['PR(>F)'][0]
                significance = 'Significant' if anova_pvalue < 0.05 else 'Not Significant'
                anova_results.append({
                    'Column': intcol,
                    'Association': significance,
                    'p-value': round(anova_pvalue, 3)
                })

            # Plots
            for col in numerical_cols:
                fig = px.histogram(
                    df_category,
                    x=col,
                    color=target_column,
                    nbins=30,
                    title=f'Histogram of {col} by {target_column}',
                    template='plotly_white',
                )
                fig.update_layout(title_x=0.5)  # Center-align the title
                plots.append(fig)

            # Correlation Matrices
            if len(numerical_cols) > 1:
                corr_methods = ['pearson', 'spearman', 'kendall']
                color_scales = [px.colors.sequential.Greys, px.colors.sequential.Blues, px.colors.sequential.Purples]
                for method, color_scale in zip(corr_methods, color_scales):
                    corr_matrix = df_category[numerical_cols].corr(method=method)
                    response_llm= self.chat_llm(corr_matrix)
                    fig_corr = px.imshow(
                        corr_matrix,
                        text_auto=True,
                        title=f'{method.capitalize()} Correlation Matrix',
                        color_continuous_scale=color_scale,
                        template='plotly_white'
                    )
                    fig_corr.update_layout(width=800, height=600, title_x=0.5)  # Center-align the title
                    corr_plots.append(fig_corr)
                    response_llm_correlation.append(f"For {method} correlation: {response_llm}")
                    corr_tables.append(corr_matrix)
            else:
                corr_plots = []
                response_llm_correlation= []
                corr_tables= []

        else:
            # Handle the case where target is continuous
            categorical_columns = list(df_category.select_dtypes(include=['object', 'category', 'bool']))
            # Exclude datetime columns
            categorical_columns = [col for col in categorical_columns if col not in datetime_cols]

            numerical_cols = list(df_category.select_dtypes(include=[np.number]))
            if target_column in numerical_cols:
                numerical_cols.remove(target_column)
            # Exclude datetime columns
            numerical_cols = [col for col in numerical_cols if col not in datetime_cols]

            # ANOVA Tests
            for col in categorical_columns:
                if df_category[col].nunique() < 2:
                    continue  # Skip if less than 2 categories
                model = ols(f'{target_column} ~ C({col})', data=df_category).fit()
                anova_table = sm.stats.anova_lm(model, typ=1)
                anova_pvalue = anova_table['PR(>F)'][0]
                significance = 'Significant' if anova_pvalue < 0.05 else 'Not Significant'
                anova_results.append({
                    'Column': col,
                    'Association': significance,
                    'p-value': round(anova_pvalue, 3)
                })

            # Scatter Matrix Plot
            if len(numerical_cols) > 0:
                plot_sc = px.scatter_matrix(
                    df_category,
                    dimensions=numerical_cols,
                    color=target_column,
                    title=f'Scatter Matrix with {target_column}',
                    template='plotly_white',
                    color_continuous_scale=px.colors.sequential.Blues
                )
                plot_sc.update_layout(width=800, height=600, title_x=0.5)  # Center-align the title
                plots.append(plot_sc)

            # Correlation Matrices
            if len(numerical_cols) > 1:
                numerical_cols_with_target = numerical_cols + [target_column]
                corr_methods = ['pearson', 'spearman', 'kendall']
                color_scales = [px.colors.sequential.Greys, px.colors.sequential.Blues, px.colors.sequential.Purples]
                for method, color_scale in zip(corr_methods, color_scales):
                    corr_matrix = df_category[numerical_cols_with_target].corr(method=method)
                    response_llm= self.chat_llm(corr_matrix)
                    fig_corr = px.imshow(
                        corr_matrix,
                        text_auto=True,
                        title=f"{method.capitalize()} Correlation Matrix",
                        color_continuous_scale=color_scale,
                        template='plotly_white'
                    )
                    fig_corr.update_layout(width=800, height=600, title_x=0.5)  # Center-align the title
                    corr_plots.append(fig_corr)
                    response_llm_correlation.append(f"For {method} correlation: {response_llm}")
                    corr_tables.append(corr_matrix)
            else:
                corr_plots = []
                response_llm_correlation= []
                corr_tables= []
                

            # Time series plot if datetime columns exist and target variable is continuous
            if len(self.datetime_cols) > 0:
                for dt_col in self.datetime_cols:
                    # Drop rows where datetime or target is NaN
                    df_time = self.df[[dt_col, self.target_col]].dropna()
                    # Sort by datetime
                    df_time = df_time.sort_values(by=dt_col)
                    fig = px.line(
                        df_time,
                        x=dt_col,
                        y=self.target_col,
                        title=f'Time Series Plot of {self.target_col} over {dt_col}',
                        template='plotly_white',
                    )
                    fig.update_layout(title_x=0.5)  # Center-align the title
                    time_series_figures.append(fig)

        # Create DataFrames from the results
        if chi_square_results:
            chi_square_df = pd.DataFrame(chi_square_results)
        else:
            chi_square_df = pd.DataFrame(columns=['Column', 'Association', 'p-value'])
            chi_square_df.loc[0] = ['No significant associations found or target is continuous', '', '']

        if anova_results:
            anova_df = pd.DataFrame(anova_results)
        else:
            anova_df = pd.DataFrame(columns=['Column', 'Association', 'p-value'])
            anova_df.loc[0] = ['No significant associations found', '', '']

        return chi_square_df, anova_df, plots, corr_plots, time_series_figures, response_llm_correlation, corr_tables  # Moved time_series_figures here

    def add_dataframe_to_docx(self, document, df):
        # Add a table to the document
        if df is not None and not df.empty:
            table = document.add_table(rows=(df.shape[0]+1), cols=df.shape[1])
            table.style = 'Light List Accent 1'

            # Add the header row.
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                hdr_cells[i].text = str(column)

            # Add the data rows.
            for i in range(df.shape[0]):
                row_cells = table.rows[i+1].cells
                for j in range(df.shape[1]):
                    row_cells[j].text = str(df.iat[i, j])
        else:
            document.add_paragraph("No data available.")

    def datainsights(self):
        shape_df, describe_df, dtypes_df, na_df, df_cat, outlier_df, target_figures, figures, box_figures, pie_figures, unique_df_count = self.univariate_analysis()
        chi_square_df, anova_df, plots, corr_plots, time_series_figures,  response_llm_correlation, corr_tables= self.multivariate_analysis()
        feature_importance_df = self.__decision_tree_feature_importance()
        feature_importance_df_gradient= self.__decision_tree_gradientboosting_feature_importance()
        response_llm_describe= self.chat_llm(describe_df)
        response_llm_dtypes= self.chat_llm(dtypes_df)
        response_llm_nadf= self.chat_llm(na_df)
        response_llm_dfcat= self.chat_llm(df_cat)
        response_llm_outlier= self.chat_llm(outlier_df)
        response_llm_uniquedf= self.chat_llm(unique_df_count)
        response_llm_chi_square= self.chat_llm(chi_square_df)
        response_llm_anova= self.chat_llm(anova_df)
        response_llm_randomforest= self.chat_llm(feature_importance_df)
        response_llm_decisiontree= self.chat_llm(feature_importance_df_gradient)
        response_llm_corr= response_llm_correlation

        # Create a Word document and add LLM responses
        document = Document()
        document.add_heading('EDA Insights Report', level=0)

        # Add shape of the DataFrame
        document.add_heading('Data Shape Information', level=1)
        document.add_paragraph(shape_df)

        # Add Descriptive Statistics
        document.add_heading('Descriptive Statistics', level=1)
        self.add_dataframe_to_docx(document, describe_df)
        document.add_paragraph(response_llm_describe)

        # Add Data Types
        document.add_heading('Data Types of Columns', level=1)
        dtypes_df_transposed = dtypes_df.T.reset_index().rename(columns={"index": "Column", 0: "Data Type"})
        self.add_dataframe_to_docx(document, dtypes_df_transposed)
        document.add_paragraph(response_llm_dtypes)

        # Add Unique Value Counts
        document.add_heading('Unique Value Counts', level=1)
        self.add_dataframe_to_docx(document, unique_df_count)
        document.add_paragraph(response_llm_uniquedf)

        # Add NA Values
        document.add_heading('NA Values in Data', level=1)
        na_df_transposed = na_df.T.reset_index().rename(columns={"index": "Column", 0: "Count of NA"})
        self.add_dataframe_to_docx(document, na_df_transposed)
        document.add_paragraph(response_llm_nadf)

        # Add Value Counts for Categorical Data
        if df_cat is not None:
            document.add_heading('Value Counts for Categorical Data', level=1)
            self.add_dataframe_to_docx(document, df_cat)
            document.add_paragraph(response_llm_dfcat)

        # Add Outlier Analysis
        if outlier_df is not None:
            document.add_heading('Outlier Analysis', level=1)
            self.add_dataframe_to_docx(document, outlier_df)
            document.add_paragraph(response_llm_outlier)

        # Add Chi-Square Analysis
        document.add_heading('Chi-Square Significance with Target', level=1)
        self.add_dataframe_to_docx(document, chi_square_df)
        document.add_paragraph(response_llm_chi_square)

        # Add ANOVA Analysis
        document.add_heading('ANOVA Significance with Target', level=1)
        self.add_dataframe_to_docx(document, anova_df)
        document.add_paragraph(response_llm_anova)

        # Add Correlation Analysis
        if response_llm_corr:
            document.add_heading('Correlation Analysis', level=1)
            for corr_text, corr_tab in zip(response_llm_corr, corr_tables):
                self.add_dataframe_to_docx(document, corr_tab)
                document.add_paragraph(corr_text)

        # Add Random Forest Feature Importance
        document.add_heading('Random Forest Feature Importance', level=1)
        self.add_dataframe_to_docx(document, feature_importance_df)
        document.add_paragraph(response_llm_randomforest)

        # Add Decision Tree Feature Importance
        document.add_heading('Decision Tree Feature Importance', level=1)
        self.add_dataframe_to_docx(document, feature_importance_df_gradient)
        document.add_paragraph(response_llm_decisiontree)

        # Save the document
        document.save('Data_Analysis_Report.docx')
        print("Data Analysis Report has been saved as 'Data_Analysis_Report.docx'.")

        # Generate the Datapane report as before
        report = dp.Select(blocks=[
            dp.Group(
                dp.Group(dp.Text("# DataGUI Insights"), dp.Text("### <div align='center'> Descriptive Statistics"), dp.DataTable(describe_df)),
                dp.Group(dp.Text(f"#### {shape_df}"), dp.Text("### <div align='center'> Data Types of Columns"), dp.DataTable(dtypes_df.T.reset_index().rename(columns={"index": "column", 0: "data type"}))),
                dp.Group(dp.Text("### <div align='center'> Unique Values in Data"), dp.DataTable(unique_df_count)),
                dp.Group(dp.Text("### <div align='center'> Checking for NA Values"), dp.DataTable(na_df.T.reset_index().rename(columns={"index": "column", 0: "Count of NA"}))),
                dp.Group(dp.Text("### <div align='center'> Value Counts Data"), dp.DataTable(df_cat) if df_cat is not None else dp.Text("No categorical columns to display value counts")),
                dp.Group(dp.Text("### <div align='center'> Outlier Analysis"), dp.DataTable(outlier_df)) if outlier_df is not None else dp.Text("No outlier information available"),
                # dp.Group(dp.Text("### <div align='center'> Class Imbalance"), dp.DataTable(pie_figures)),
                *[dp.Plot(fig) for fig in target_figures],
                *[dp.Plot(fig) for fig in figures],
                *[dp.Plot(fig) for fig in box_figures],
                columns=2,
                label="Univariate Analysis"
            ),
            dp.Group(
                dp.Group(dp.Text("### <div align='center'> Chi-Square Significance with Target"), dp.DataTable(chi_square_df)),
                dp.Group(dp.Text("### <div align='center'> ANOVA Significance with Target"), dp.DataTable(anova_df)),
                dp.Group(
                    dp.Text("### <div align='center'>Random Forest Feature Importance Score"),
                    dp.DataTable(feature_importance_df.reset_index(drop=True)),
                    label="ML Based Feature Importance1"
                ), dp.Group(
                    dp.Text("### <div align='center'>Decision Tree Feature Importance Score"),
                    dp.DataTable(feature_importance_df_gradient.reset_index(drop=True)),
                    label="ML Based Feature Importance2"
                ),
                *[dp.Plot(fig) for fig in corr_plots] if corr_plots else dp.Text("Correlation matrices not generated due to insufficient numerical columns."),
                *[dp.Plot(fig) for fig in plots],
                *[dp.Plot(fig) for fig in time_series_figures],  # Added time_series_figures here
                columns=2,
                label="Multivariate Analysis"
            ),

        ])
        dp.save_report(report,"testing.html", open=True, formatting=dp.Formatting(
                        width=dp.Width.FULL,
                        text_alignment=dp.TextAlignment.CENTER,
                        light_prose=False,
                        font=dp.FontChoice.SERIF, 
                    ))

        return report

# Example usage:
# df = sns.load_dataset("tips")
# DataAnalyzeObject = AnalyzeData(df, "tip")
# DataAnalyzeObject.datainsights()