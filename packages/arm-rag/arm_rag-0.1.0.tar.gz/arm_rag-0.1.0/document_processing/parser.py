import os
import yaml
import pypdf
import docx.api
import spire.doc
import zipfile
import tempfile
from config import CONFIG

class Document_parser:
    def __init__(self):
        self.supported_formats = CONFIG['document']['supported_formats']
    
    def parse(self, path):
        """The main method which parses initial file/archive  from the user"""
        
        if path.endswith('.zip'):
            return self.parse_archive(path) # returns dict of multiple documents
        else:
            content = self.parse_file(path)
            head, filename = os.path.split(path)
            return {filename: content} # returns dict of signle document

    def parse_archive(self, archive_path):
        """returns a dict: {filename: content}"""
        file_contents = {}
        with zipfile.ZipFile(archive_path, 'r') as zip_ref:
            with tempfile.TemporaryDirectory() as temp_dir:
                for member in zip_ref.namelist():
                    extracted_path = zip_ref.extract(member, path=temp_dir)
                    content = self.parse_file(extracted_path)
                    file_contents[member] = content
        return file_contents

    def parse_file(self, file_path):
        if '.' + file_path.split('.')[-1] in self.supported_formats:
            if file_path.endswith('.pdf'):
                content = self.read_pdf(file_path)
            elif file_path.endswith('.txt'):
                content = self.read_txt(file_path)
            elif file_path.endswith('.doc'):
                content = self.read_doc(file_path)
            elif file_path.endswith('.docx'):
                content = self.read_docx(file_path)
            
            return content
        else:
            return None #TODO: change None return with smth else
    
    def read_pdf(self, path):
        with open(path, 'rb') as file:
            reader = pypdf.PdfReader(file, strict=False)
            text = ""
            for page in reader.pages:
                content = page.extract_text()
                text += content
        return text
    
    def read_txt(self, path):
        text = ""
        with open(path, 'r') as file:
            for line in file:
                text += line 
        return text
    
    def read_docx(self, path):
        text = ""
        doc = docx.api.Document(path)
        for p in doc.paragraphs:
            text += p.text + "\n"
        return text

    def read_doc(self, path):
        document = spire.doc.Document()
        document.LoadFromFile(path)
        section = document.Sections[0]
        text = ''
        for i in range(section.Paragraphs.Count):
            paragraph = section.Paragraphs[i]
            text += paragraph.Text + "\n"
        return text
    
############### Test

# config_path = os.path.join(os.path.dirname(__file__), "../config.yaml")
# with open(config_path, 'r') as f:
#     config = yaml.safe_load(f)
    
# parser = Document_parser(config)

# # single pdf test
# # text = parser.parse(r'C:\Users\Lenovo\Desktop\arm_doc_chat\uploaded_files\a173a22ad75a77ce78174ae6953089e2.pdf')
# # print(text)

# # zip file test
# content_dict = parser.parse(r'C:\Users\Lenovo\Desktop\arm_doc_chat\uploaded_files\matanaliz_xndragirq.zip')
# print(content_dict)
# print(type(content_dict))
# print(content_dict.keys())
