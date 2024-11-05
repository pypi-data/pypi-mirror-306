# -*- coding:UTF-8 -*-
import json
import re
import traceback
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from loguru import logger

# noinspection PyProtectedMember


class DocxUtil(object):
    def __init__(self, docPath):
        # docPath路径必须存在,该功能中不会自己创建目录
        self.basePath = docPath

    def creatDocx(self, jsonStr, link_dicts=None):
        """
        根据json字符串生成docx测试报告,json.loads默认无法解析单引号字符串,需处理成双引号
        :param jsonStr: json字符串
        :return:
        """
        try:
            ty = type(jsonStr).__name__
            if ty != "dict":
                if ty != "str":
                    logger.error("传入的json必须为字符串!")
                    return False

                try:
                    dict = json.loads(jsonStr)
                except Exception as ex:
                    logger.error("\tError %s\n" % ex)
                    logger.error("json格式错误,单引号数据请自行解析可用replace()或eval()等函数")
                    return False
            else:

                dict = jsonStr

            document = Document()

            # 设置整个文档的默认字体
            document.styles['Normal'].font.name = '微软雅黑'
            document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            for key in dict.keys():
                if key == "title":
                    # print(key + "-标题:", dict[key])
                    # 设置文档标题的字体
                    titleStyle = document.add_heading("", 0)

                    # 主标题居中
                    titleStyle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    titleFont = titleStyle.add_run(dict[key])
                    titleFont.font.name = "微软雅黑"
                    titleFont._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


                else:
                    # 一级标题样式
                    run = document.add_heading('', level=1)
                    # 段前段后
                    run.paragraph_format.space_before = Pt(1.5)
                    run.paragraph_format.space_after = Pt(1.5)
                    # 行间距
                    run.paragraph_format.line_spacing = 1.5
                    # 左缩进,右缩进 英寸
                    run.paragraph_format.left_indent = Inches(0)
                    run.paragraph_format.right_indent = Inches(0)

                    # 标题
                    fontStyle = run.add_run(key)
                    # 标题字体
                    fontStyle.font.name = "微软雅黑"
                    fontStyle._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    # 字体大小,小四 为12磅
                    fontStyle.font.size = Pt(12)
                    # 标题颜色,不设置为蓝色
                    fontStyle.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色

                    # 创建表格
                    table = None
                    # document.add_paragraph("这是关于"+key+"的描述内容", style='List Number')
                    # document.add_paragraph("这是关于" + key + "的描述内容", style='List Bullet')
                    if 'dict' in str(type(dict[key])):
                        # 段落描述字体
                        # table = document.add_table(rows=len(dict[key]), cols=2, style='Table Grid')
                        table = document.add_table(rows=len(dict[key]), cols=2, style='Medium Grid 1 Accent 4')
                        # 设置单元格列宽
                        table.cell(0, 0).width = Cm(3)
                        table.cell(0, 1).width = Cm(15)
                        index = 0
                        for keys_dict in dict[key].keys():

                            if 'list' in str(type(dict[key][keys_dict])):

                                i = 0
                                if len(dict[key][keys_dict]) != 0:
                                    for k in dict[key][keys_dict]:
                                        if i != 0:
                                            table.add_row()
                                        table.rows[index].cells[0].text = keys_dict

                                        n = re.findall(r"<color=red>(.+?)</color>", k)
                                        if n is not None and len(n) > 0:
                                            table.rows[index].cells[1].text = ""
                                            p = table.rows[index].cells[1].paragraphs[0]
                                            for text in n:
                                                k = k.replace("<color=red>" + text + "</color>", "@&*^%@")

                                            strs = k.split("@&*^%@")
                                            for i, s in enumerate(strs):
                                                run_black = p.add_run(s)
                                                run_black.font.color.rgb = RGBColor(0, 0, 0)
                                                run_black.font.bold = False
                                                if i < len(n):
                                                    run_red = p.add_run(n[i])
                                                    run_red.font.color.rgb = RGBColor(255, 0, 0)
                                                    run_red.font.bold = False


                                        else:
                                            # table.rows[index].cells[1].text = k
                                            # todo 判断是否需要添加链接(第三部分数据)
                                            link_txt = None
                                            for link_key in link_dicts.keys():
                                                if link_key in k:
                                                    link_txt = link_key
                                                    break

                                            if link_txt is not None and link_txt in k:
                                                p = table.rows[index].cells[1].paragraphs[0]
                                                self.add_hyperlink(p, link_txt, link_dicts[link_txt])
                                                if link_txt != k:
                                                    self.append_text(p, k.replace(link_txt, ""))
                                            else:
                                                table.rows[index].cells[1].text = k

                                        # print(str(i),keys_dict,str(index),str(len(dict[key][keys_dict])))

                                        # 合并左侧标题单元格
                                        if 1 < len(dict[key][keys_dict]) != (i + 1):
                                            table.cell(index, 0).merge(table.cell(index + +1, 0))
                                            # 设置列宽
                                            # table.cell(0, 0).width = Cm(3)
                                            # table.cell(0, 1).width = Cm(15)
                                        # if index - len(dict[key][keys_dict]) < 2:
                                        #     table.cell(index, 0).merge(table.cell(index + 1, 0))

                                        # 去除默认第一行加粗样式
                                        run_0 = table.rows[index].cells[0].paragraphs[0].runs[0]
                                        run_1 = table.rows[index].cells[1].paragraphs[0].runs[0]
                                        run_0.font.bold = False
                                        run_1.font.bold = False
                                        # 去除默认第一行加粗样式

                                        index = index + 1
                                        i = i + 1
                                else:
                                    table.add_row()
                                    table.rows[index].cells[0].text = keys_dict
                                    k = dict[key][keys_dict]

                                    n = re.findall(r"<color=red>(.+?)</color>", k)
                                    if n is not None and len(n) > 0:
                                        table.rows[index].cells[1].text = ""
                                        p = table.rows[index].cells[1].paragraphs[0]
                                        for text in n:
                                            k = k.replace("<color=red>" + text + "</color>", "@&*^%@")

                                        strs = k.split("@&*^%@")
                                        for i, s in enumerate(strs):
                                            run_black = p.add_run(s)
                                            run_black.font.color.rgb = RGBColor(0, 0, 0)
                                            run_black.font.bold = False
                                            if i < len(n):
                                                run_red = p.add_run(n[i])
                                                run_red.font.color.rgb = RGBColor(255, 0, 0)
                                                run_red.font.bold = False
                                    else:

                                        table.rows[index].cells[1].text = dict[key][keys_dict]

                                    # 去除默认第一行加粗样式
                                    run_0 = table.rows[index].cells[0].paragraphs[0].runs[0]
                                    run_1 = table.rows[index].cells[1].paragraphs[0].runs[0]
                                    run_0.font.bold = False
                                    run_1.font.bold = False
                                    # 去除默认第一行加粗样式

                                    index = index + 1

                            else:
                                table.rows[index].cells[0].text = keys_dict
                                k = dict[key][keys_dict]

                                n = re.findall(r"<color=red>(.+?)</color>", k)
                                if n is not None and len(n) > 0:
                                    # print("-------------"+k+"------|")
                                    table.rows[index].cells[1].text = ""
                                    p = table.rows[index].cells[1].paragraphs[0]
                                    for text in n:
                                        k = k.replace("<color=red>" + text + "</color>", "@&*^%@")

                                    strs = k.split("@&*^%@")
                                    for i, s in enumerate(strs):
                                        # wukeke===========
                                        # link_txt = None
                                        # for link_key in link_dicts.keys():
                                        #     if link_key in s:
                                        #         link_txt = link_key
                                        #         break
                                        #
                                        # if link_txt is not None and link_txt in s:
                                        #     self.add_hyperlink(p, link_txt, s)
                                        #     if link_txt != s:
                                        #         self.append_text(p, s.replace(link_txt, ""))
                                        # else:
                                        # wukeke===========

                                        # self.add_hyperlink(p, s, "www.baidu.com")

                                        run_black = p.add_run(s)
                                        run_black.font.color.rgb = RGBColor(0, 0, 0)
                                        run_black.font.bold = False
                                        if i < len(n):
                                            run_red = p.add_run(n[i])
                                            run_red.font.color.rgb = RGBColor(255, 0, 0)
                                            run_red.font.bold = False
                                else:

                                    # table.rows[index].cells[1].text = dict[key][keys_dict]
                                    # todo 判断是否需要添加链接 加超链接判断(第一部分数据)
                                    link_txt = None
                                    for link_key in link_dicts.keys():
                                        if link_key in k:
                                            link_txt = link_key
                                            break

                                    if link_txt is not None and link_txt in dict[key][keys_dict]:
                                        p = table.rows[index].cells[1].paragraphs[0]
                                        self.add_hyperlink(p, link_txt, link_dicts[link_txt])
                                        if link_txt != dict[key][keys_dict]:
                                            self.append_text(p, dict[key][keys_dict].replace(link_txt, ""))
                                    else:
                                        table.rows[index].cells[1].text = dict[key][keys_dict]

                                # 去除默认第一行加粗样式
                                run_0 = table.rows[index].cells[0].paragraphs[0].runs[0]
                                run_1 = table.rows[index].cells[1].paragraphs[0].runs[0]
                                run_0.font.bold = False
                                run_1.font.bold = False
                                # 去除默认第一行加粗样式

                                index = index + 1

                        # for row, obj_row in enumerate(table.rows):
                        #     for col, cell in enumerate(obj_row.cells):
                        #         cell.text = cell.text + "%d,%d " % (row, col)


                    elif 'list' in str(type(dict[key])):
                        if len(dict[key]) != 0:
                            table = document.add_table(rows=(len(dict[key])) + 1, cols=len(dict[key][0]),
                                                       style='Medium Grid 1 Accent 5')

                            hdr_cells = table.rows[0].cells
                            for i, key_head in enumerate(dict[key][0]):
                                hdr_cells[i].text = key_head

                            for i, map in enumerate(dict[key]):
                                for j, key2 in enumerate(map.keys()):
                                    table.rows[i + 1].cells[j].text = map[key2]

                                    # 去除默认样式第一行加粗
                                    run = table.rows[i + 1].cells[j].paragraphs[0].runs[0]
                                    run.font.bold = False
                                    # 去除默认样式第一行加粗

            try:
                document.save(self.basePath)

            except Exception as ex:
                if "No such file or directory" in str(ex):
                    logger.error("目录不存在,请检查,\tError %s\n" % ex)
                else:
                    logger.error("\tError %s\n" % ex)
                return False

        except Exception as ex:
            logger.error("\tError %s\n" % ex)
            logger.error(traceback.format_exc())
            return False
        return True

    def add_hyperlink(self, paragraph, text, url):
        """
        function used for
        :__Author__ : wu.keke
        :__mtime__ : 2020/11/19 10:02
        """
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        r = paragraph.add_run()
        r._r.append(hyperlink)

        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return new_run

    def append_text(self, paragraph, text):
        """
        function used for
        :__Author__ : wu.keke
        :__mtime__ : 2020/11/19 10:02
        """

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        r = paragraph.add_run()
        r._r.append(hyperlink)

        # r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = False

    def add_link(self, text, url):
        document = Document()
        p = document.add_paragraph('段落')
        self.add_hyperlink(p, '百度', 'www.baudu.com')
        document.save(r"D:\2020\2020持仓检查与比对\demo_link.docx")
        return True


if __name__ == '__main__':
    # jsonStr = "{\"title\": \"准入测试报告\", \"一、准入测试报告\": {\"系统\": \"NADDONTEST\", \"版本\": \"交付件检查2\", \"时间\": \"2020-05-07 17:20:33\", \"补丁\": [], \"结果\": \"不通过\"}, \"二、准入测试版本发现问题清单\": [{\"编号\": \"1\", \"分类\": \"版本交付单_该版本下无任何开发任务或问题单\", \"描述\": \" 无效交付\"}, {\"编号\": \"2\", \"分类\": \"版本交付单_交付类文档缺失文件名包含关联表的文档\", \"描述\": \"  缺失必须交付的关联表文档\"}, {\"编号\": \"3\", \"分类\": \"版本交付单_交付类文档缺失文件名包含影响分析的文档\", \"描述\": \"  缺失必须交付的影响分析文档\"}, {\"编号\": \"4\", \"分类\": \"版本交付单_交付类文档缺失文件名包含升级or用户的文档\", \"描述\": \"  缺失必须交付的升级or用户文档\"}], \"三、准入测试交付路径\": {\"交付单\": \"NADDONTEST-594 不触发交付件检查单2\", \"交付路径\": \"http://artifactory.test.com:8081/artifactory/Delivery/NADDONTEST/交付件检查2/8/\"}}"
    # jsonStr = "{\"title\": \"准入测试报告\", \"一、准入测试报告\": {\"系统\": \"NADDONTEST\", \"版本\": \"交付件检查1\", \"时间\": \"2020-05-14 21:11:13\", \"补丁\": [\"NADDONTEST-636交付件检查1的Sub-开发子任务\", \"NADDONTEST-585交付件检查1开发任务\", \"NADDONTEST-41开发任务测试\"], \"结果\": \"通过\"}, \"二、准入测试版本发现问题清单\": [], \"三、准入测试交付路径\": {\"交付单\": \"NADDONTEST-594 不触发交付件检查单2\", \"交付路径\": \"http://artifactory.test.com:8081/artifactory/Delivery/NADDONTEST/交付件检查1/115/\"}}"
    # jsonStr = "{\"title\": \"准入测试报告\", \"一、准入测试报告\": {\"系统\": \"NADDONTEST\", \"版本\": \"交付件检查2\", \"时间\": \"2020-05-25 09:52:25\", \"结果\": \"通过\"}, \"二、补丁列表\": {\"NADDONTEST-636\": \"交付件检查1的Sub-开发子任务\", \"NADDONTEST-543\": \"test\", \"NADDONTEST-535\": \"wd\"}, \"三、准入测试交付路径\": {\"交付单\": \"NADDONTEST-584 不触发交付件检查单1\", \"交付路径\": \"http://artifactory.test.com:8081/artifactory/Delivery/NADDONTEST/交付件检查2/25/\"}, \"四、文档追加记录\": {\"追加文档路径\": \"http://artifactory.test.com:8081/artifactory/Delivery/NADDONTEST/交付件检查2/25/文档_20200525095126/\"}}"
    # jsonStr = ""
    # DocxUtil("D:\\2020\\mydemo2.docx").creatDocx(jsonStr)
    dicts = {'title': 'NADDONTEST_交付件检查2_72_准入测试报告_通过',
             '一、准入测试简报': {'系统': 'NADDONTEST', '版本': '交付件检查2', '批次号': '72', '准入服务模式': '全量交付',
                          '准入服务开始时间': '2020-06-11 16:22:52', '准入服务结束时间': '2020-06-15 11:12:39',
                          '准入服务结果': '<color=red>通过</color>',
                          '准入交付测试路径': 'http://artifactory.test.com:8081/artifactory/Delivery/NADDONTEST/交付件检查2/72/'},
             '二、准入服务风险提示': {
                 '自动交付件检查模块': '[1]版本交付单_交付类文档缺失文件名包含FS的文档:  <color=red>缺失必须</color>交付的FS文档\n[2]版本交付单_交付类文档缺失文件名包含IS的文档:  缺失必须交付的IS文档\n[3]版本交付单_交付类文档缺失文件名包含保障的文档:  缺失必须交付的保障文档\n[4]版本交付单_交付类<color=red>文档缺失</color>文件名包含应急or流程卷的文档:  缺失必须交付的应急or流程卷文档\n[5]版本交付单_交付类文档缺失文件名包含运维的文档:  缺失必须交付的运维文档\n[6]版本交付单_交付类文档缺失文件名包含应用支持or工具的文档:  缺失必须交付的应用支持or工具文档\n'},
             '三、版本JIRA单统计': {
                 '版本交付单': ['NADDONTEST-584 不触发交付件检查单1-"双引号"-\'单引号\'-`特殊符号` 准入测试中', 'NADDONTEST-599 报错问题还原单 已关闭'],
                 '开发任务': '无', 'Sub-开发任务': ['NADDONTEST-636 交付件检查1的Sub-开发子任务 系统测试中'],
                 '需求': ['NADDONTEST-616 交付件检查用的测试需求单 开始'],
                 '系统测试问题单': ['NADDONTEST-544 <color=red>测试关闭</color>问题单1号<color=red>红色内容</color> 已关闭',
                             'NADDONTEST-543 test 验证通过',
                             'NADDONTEST-535 【业务管理平台】通知通告新建接口和准入规则详情页接口报"“Request method \'GET\' not supported”" 系统测试中'],
                 '安全问题单': '无', '生产问题单': '无',
                 '准入问题单': ['NADDONTEST-696 交付件检查2-第53次交付-交付件检查不通过 开始', 'NADDONTEST-695 交付件检查2-第53次交付-交付件检查不通过 开始',
                           'NADDONTEST-694 交付件检查2-第53次交付-交付件检查不通过 开始', 'NADDONTEST-693 交付件检查2-第53次交付-交付件检查不通过 开始',
                           'NADDONTEST-686 交付件检查2-第43次交付-交付件检查不通过 已完成', 'NADDONTEST-676 交付件检查2-第39次交付-交付件检查不通过 已完成',
                           'NADDONTEST-660 交付件检查2-第6次交付-交付件检查不通过 已完成', 'NADDONTEST-659 交付件检查2-第6次交付-交付件检查不通过 已完成',
                           'NADDONTEST-658 交付件检查2-第6次交付-交付件检查不通过 已完成', 'NADDONTEST-656 交付件检查2-第6次交付-交付件检查不通过 已完成',
                           'NADDONTEST-655 交付件检查2-第6次交付-交付件检查不通过 已完成', 'NADDONTEST-654 交付件检查2-第6次交付-交付件检查不通过 已完成',
                           'NADDONTEST-653 交付件检查2-第6次交付-交付件检查不通过 已完成', 'NADDONTEST-652 交付件检查2-第5次交付-交付件检查不通过 已完成',
                           'NADDONTEST-639 交付件检查2-第1次交付-交付件检查不通过 已完成', 'NADDONTEST-638 交付件检查2-第1次交付-交付件检查不通过 已完成',
                           'NADDONTEST-634 交付件检查2-第1次交付-交付件检查不通过 已完成', 'NADDONTEST-595 交付件检查2-第1次交付-交付件检查不通过 已完成'],
                 '交付问题单': '<color=red>无</color>'},
             '四、多列测试': [{'一': 'hhhh','二':'2222','三':'33333'}, {'二2:': 'uuuuu','哈哈':'jjjjjj'}, {'三2': 'yyyyy'}]
             }

    link_dicts = {"全量": "http://eqops.tc.com/confluence/display/SCAD/2.ETF",
                  "NADDONTEST-599": "http://eqops.tc.com/confluence/display/SCAD/2.ETF", "流程卷的文档:": "www.baidu.com"}
    DocxUtil("./mydemo5.docx").creatDocx(dicts, link_dicts)
