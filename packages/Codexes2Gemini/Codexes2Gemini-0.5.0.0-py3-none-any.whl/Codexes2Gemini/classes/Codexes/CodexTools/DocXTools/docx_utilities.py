#  Copyright (c) 2023. Fred Zimmerman.  Personal or educational use only.  All commercial and enterprise use must be licensed, contact wfz@nimblebooks.com

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# this holds various docx utilities that are too complex to include in a pipeline widget

# the first three allows you to add a figure, table, or index entry to a paragraph

# usage:
# document = Document()
# paragraph = document.add_paragraph('')

# paragraph = document.add_paragraph('Figure ', style='Caption')
# Figure(paragraph) <==
# paragraph.add_run(' Figure Caption ')
# #document.add_picture(FileDir + FileName, width=Cm(10))
# paragraph = document.add_paragraph('Figure Text', style='BodyText')
# document.add_page_break()

def MarkIndexEntry(entry, paragraph):
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' XE "%s" ' % (entry)
    r.append(instrText)

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def Figure(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Figure \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def Table(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
