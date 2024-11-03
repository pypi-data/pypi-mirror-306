#  Copyright (c) 2023. Fred Zimmerman.  Personal or educational use only.  All commercial and enterprise use must be licensed, contact wfz@nimblebooks.com

import argparse

import streamlit as st
from docx import Document


def setup(df_filename, docx_filename):  # , basedocx):

    # if basedocx:
    #     document = Document(basedocx)
    # else:
    #     document = Document()

    # document.save(os.path.join(output_dir, docx_filename))

    st.info('docx2df2changelist2docx.py')

    doc = Document(docx_filename)

    df = df_filename
    # doc = docx_filename
    df = df.loc[~df['table_id'].str.isdigit()]
    df.shape
    # reset index
    df = df.reset_index(drop=True)
    print('number of items in df omitting table rows' + str(len(df.index)))
    # count the number of rows in the dataframe
    num_rows = len(df.index)

    # count the number of paragraphs in the document
    number_paragraphs = len(doc.paragraphs)

    if num_rows != number_paragraphs:
        print('number of rows in df {}'.format(num_rows) + 'does not match number of paragraphs in doc {}'.format(
            number_paragraphs))
        print(
            'Automatic editing may be unreliable.  Make sure the rows in the dataframe match the paragraphs in the document.')
        exit()
    else:
        print('Number of rows in dataframe matches number of paragraphs in document')

    return df, doc


def run_fonts_check_and_change(runs):
    for r in runs:
        r.font.name = None
        r.font.size = None
    return runs


def para_with_direct_formatting2cleanpara(p):
    p.paragraph_format.line_spacing = None
    p.paragraph_format.left_indent = None
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = None
    p.paragraph_format.widow_control = True
    return p


def parse_requested_style(requested_style):
    try:
        actions_dict = {'style2bodytext': 'Body Text', 'style2caption': 'Caption', 'style2bibliography': 'Bibliography',
                        'style2byline': 'Byline', 'style2title': 'Title', 'style2normal': 'Normal',
                        'style2quote': 'Quote', 'style2quotesignature': 'quote signature',
                        'style2heading1': 'Heading 1', 'style2part': 'Part', 'style2signature': 'Signature'}
        return actions_dict[requested_style]
    except Exception as e:
        print('error parsing requested style')
        print(e)
        return e


def get_requested_action(df, doc, rows):
    styles = doc.styles  # get the styles collection
    print(styles['Body Text'].font.name)

    # drop rows where table_id is not null
    print('initial length of table: {}'.format(len(df.index)))

    for i in range(len(df.index)):

        para_text = df.iloc[i]['para_text']
        requested_action = df.iloc[i]['requested_action']
        next_para_text = df.iloc[i]['para_text']
        print(i, requested_action)
        # make the change to the corresponding paragraph in the Word document
        for count, p in enumerate(doc.paragraphs):
            if p.text == para_text:
                runs = p.runs
                if isinstance(requested_action, str):
                    if requested_action.startswith("style2"):
                        try:
                            requested_style = parse_requested_style(requested_action)
                            p = para_with_direct_formatting2cleanpara(p)  # remove the font and size from the paragraph
                            run_fonts_check_and_change(runs)
                            p.style = styles[requested_style]

                            print('applying style {} to paragraph {} that begins with {}'.format(requested_style, count,
                                                                                                 para_text[0:20]))
                        except Exception as e:
                            print('could not parse requested style')
                        break
                    elif requested_action == 'make_empty_str':
                        print('making paragraph {} empty'.format(count))
                        p.text = ''
                        break
                    elif requested_action == None:
                        print('no requested action for paragraph {}'.format(count))
                    # now check if the paragraph is an image followed by a blank line that does not have a caption
                if para_text.startswith('Document_Image') and next_para_text == '':
                    df.at[i + 1, 'style'] = styles['Caption']
                    print(
                        'applying style Caption to paragraph {} that follows Document_Image in previous paragraph {}'.format(
                            count + 1, count))

        if i >= rows:
            break
    return df, doc


def argparse_handler():
    argparser = argparse.ArgumentParser()

    argparser.add_argument('--docx_filename', help='path to docx file', default='test/docx/lorem.docx')
    argparser.add_argument("--basedocx", help="base Word file",
                           default="/Users/fred/bin/nimble/unity/editorial/template.dotx")
    argparser.add_argument('--output_dir', help='path to output directory', default='output')
    argparser.add_argument('--df_filename', help='output converted text as single string, not a list',
                           default='test/csv/lorem2.csv')
    argparser.add_argument('--rows', help='number of rows to process', default=20)
    args = argparser.parse_args()
    docx_filename = args.docx_filename
    output_dir = args.output_dir
    df_filename = args.df_filename
    rows = int(args.rows)
    basedocx = args.basedocx

    return docx_filename, output_dir, df_filename, rows


if __name__ == "__main__":

    docx_file, output_dir, df_filename, rows = argparse_handler()  # add basedocx later
    print(docx_file, output_dir, df_filename)
    df, doc = setup(df_filename, docx_file)  # , basedocx)
    results = get_requested_action(df, doc, rows)
    doc = results[1]
    try:
        doc.save(output_dir + '/testing.docx')
        print('saved current_changed.docx')
    except Exception as e:
        print('problem saving file  {}'.format(e))
