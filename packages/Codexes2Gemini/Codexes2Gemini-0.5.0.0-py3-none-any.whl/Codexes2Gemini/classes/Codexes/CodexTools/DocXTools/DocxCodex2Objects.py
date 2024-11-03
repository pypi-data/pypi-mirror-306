#  Copyright (c) 2023. Fred Zimmerman.  Personal or educational use only.  All commercial and enterprise use must be licensed, contact wfz@nimblebooks.com

'''
This is a class that simplifies, streamlines, and regularizes the process of transforming a digital book codex into objects that can be read and acted upon by NLP and LLM services.

A digital codex is bytes in the shape of a book: i.e. a digital format object, such as PDF, docx, text, or markdown, that is organized in the structure of a book and is roughly the same length as a book.

This library is intended for use on "real" digital books, not on short documents or articles or text data dumps, but it is able to fall back to convert short or malformed documents to unstructured text.

The library makes use of several programs I have already written, including:
 - docx2textfile
 - docx2text
 - docx2jsonl
 - docx2chunks
 - docx2utilities
 - docx2dataframe


'''

import argparse
import os
import subprocess

import docx
import pandas as pd

from app.utilities.gpt3complete import chatcomplete
from classes.Codexes.docx2dataframe.main_code import Para_table_image_extraction as pti


class DocxCodex2Objects:
    def docxfilepaths2dfs(self, docxfilepaths, output_dir, output_format, paras_limit):
        # check if docxfiles is a list
        if isinstance(docxfilepaths, list):
            results = []
            for docxfilepath in docxfilepaths:
                # check if docxfile is a valid docx
                docxvalidation = self.docxvalidate(docxfilepath)
                print(docxvalidation)
                if docxvalidation:
                    print('docx file is valid')
                    print(f"converting {docxfilepath} to tuple of dataframes")
                    # convert docx to objects
                    combined_df, image_df, list_of_tables = pti.main(docxfilepath)
                    combined_df['docxfilepath'] = docxfilepath
                    # if image_df is not empty:
                    if not image_df.empty:
                        image_df['docxfilepath'] = docxfilepath
                    list_of_tables['docxfilepath'] = docxfilepath
                    thisfile_results = [combined_df, image_df, list_of_tables]

                    results.append(thisfile_results)
                # print(results)
            return results
        else:  # we know docxfilepaths is not a list
            docxsinglefilepath = docxfilepaths
            # check if single docxfile is a valid docx
            docxvalidation = self.docxvalidate(docxsinglefilepath)
            if docxvalidation:
                combined_df, image_df, list_of_tables = pti.main(docxsinglefilepath)
                results = combined_df, image_df, list_of_tables
                return results
            # results is list of tuple of dataframes: combined_df, image_df, list_of_tables

    def docxdfs2presets_with_prompts(self, docxdfs, prompts, presets, rows=10, beginning_row=0, ending_row=10,
                                     engine="gpt-3.5-turbo"):
        print(docxdfs)
        if isinstance(docxdfs, list):
            print('docxdfs is a list')
            results = []
        for df in docxdfs:
            if isinstance(df, pd.DataFrame):
                print('df is a dataframe')
                print('iterating through df rows')
                for index, row in df.iterrows():
                    text = row['para_text']
                    print(text)
                    # send text to presets
                    for preset in presets:
                        if not isinstance(prompts, list):
                            prompts = [prompts]
                        else:
                            for prompt in prompts:
                                try:
                                    response_text = chatcomplete(preset, prompt + text, engine)
                                    result = response_text
                                    print(result)
                                    # add result to current row of df
                                    df.loc[index, preset] = result
                                except Exception as e:
                                    print(e)
                                    pass
            print(f"results for ths dataframe \n\n{df.head(5)}))")
        return docxdfs

    def docxvalidate(self, docxfile):
        # validate file is valid docx
        try:
            docxfile = docx.Document(docxfile)
        except Exception as e:
            print(f'docx file is not valid {e}')
            return False
        return True

    def doc2docx(self, docfile, output_dir):

        # Define the output path (replace .doc with .docx)
        docx_path = output_dir.replace('.doc', '.docx')

        # Use the LibreOffice command line tools to convert
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'docx',
            '--outdir', os.path.dirname(output_dir),
            docfile
        ]

        subprocess.run(cmd)

        # Check if the conversion was successful
        if os.path.exists(docx_path):
            return docx_path
        else:
            return None

    '''
    The following function puts a document into formats that can be accepted by the Rating Utilities functions in the Synthetic Reader class.
    The "rate_ideas" function accepts a dataframe of ideas and a reader panel dataframe. The ideas dataframe must contain the key 'idea' and the values for that key must be strings of text. The tokenized values must fit within the max_tokens parameter of the LLM API.
    '''

    def prepare4readerpanel(self, docxfile):
        docxvalidation = self.docxvalidate(docxfile)
        if docxvalidation:
            doc = docx.Document(docxfile)
            # use python-docx to get the text
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            # TODO: extract ideas using llama-text
            # TODO: add ideas to dataframe
            paras = fullText
        return paras

    def send2readerpanel(self, docxfile):

        return

    def argparse_handler():
        argparser = argparse.ArgumentParser(description='Convert docx file to text, jsonl, dataframe, or chunks')

        argparser.add_argument('--docxfile', '-D', help='path to docx file', default='test/docx/lorem.docx')
        argparser.add_argument('--input_file_list', '-I', help='list of input files', default=['test/docx/lorem.docx'])
        argparser.add_argument('--output_dir', help='path to output directory', default='output')
        argparser.add_argument('--output_format',
                               '-O', help='output format: text, jsonl, dataframe, or chunks', default='text')
        argparser.add_argument('--paras_limit', help='limit number of paragraphs displayed on return', default=20)

        args = argparser.parse_args()
        # return all args as dict of argument name and value
        return vars(args)
