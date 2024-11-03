#  Copyright (c) 2023. Fred Zimmerman.  Personal or educational use only.  All commercial and enterprise use must be licensed, contact wfz@nimblebooks.com

import re

try:
    from pypandoc import convert_file
except ImportError as e:
    print("Pandoc not installed. No pypandoc module")
    print(e)
    exit(1)

from docx import Document

from docx.shared import Inches


def docxdf2docx(docx_production_df, thisdoc_dir):
    # create the docx
    document = Document("resources/docx/FrontMatter2.docx")
    try:
        for index, row in docx_production_df.iterrows():
            print(row)
            if row["style"] == "Page Break":
                document.add_page_break()
                continue
            if row["para_text"] == "":
                continue

            # remove newline if it occurs at beginning of string
            row["para_text"] = re.sub(r'^\n', '', row["para_text"])
            # if a newline is surrouonded only by spaces,
            # remove the spaces
            row["para_text"] = re.sub(r'\n\s*\n', '\n\n', row["para_text"])
            # if two or more newlines are found in a row, convert them to only one
            row["para_text"] = re.sub(r'\n{2,}', '\n', row["para_text"])
            # if a newline is found at the end of a string, remove it
            row["para_text"] = re.sub(r'\n$', '', row["para_text"])
            if row['para_text'].startswith('Scene '):
                row['style'] = 'Heading 3'
                print('found scene heading')

            if row["style"] == "Image":
                if row["para_text"] != "Image Placeholder":
                    try:
                        document.add_picture(row["para_text"], width=Inches(5.5))
                    except Exception as eimage:
                        errmessage = str(eimage) + 'error in adding image'
                        print(errmessage)
                        document.add_paragraph("Image not found, check log to see if it was generated.", row["style"])
                else:
                    document.add_paragraph(row["para_text"], row["style"])
            # if cell begins with the string "Style ", then it is a heading 3
            else:
                # count newlines in the text-bearing string
                newline_count = row["para_text"].count('\n')
                # now if there are multiple paragraphs in the string, split them and add them as separate paragraphs
                if newline_count >= 1:
                    paragraphs = row["para_text"].split('\n')
                    for p in paragraphs:
                        q = "Problem is here"
                        document.add_paragraph(p, row["style"])
                else:
                    document.add_paragraph(row["para_text"], row["style"])

    except Exception as e:
        errmessage = str(e) + 'error in creating word doc'
        print(errmessage)

    document.save(thisdoc_dir)

    # convert the doc to pdf
    try:
        convert_file(thisdoc_dir + '/' + "frontmatter.docx", 'pdf',
                     outputfile=thisdoc_dir + '/' + "frontmatter.pdf")
    except Exception as e:
        errmessage = str(e) + ' error in creating interior postscript'
        print(errmessage)

    return docx_production_df, document
