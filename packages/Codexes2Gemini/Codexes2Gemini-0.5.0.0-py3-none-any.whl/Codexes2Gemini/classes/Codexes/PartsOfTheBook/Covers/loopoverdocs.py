#  Copyright (c) 2024. Fred Zimmerman.  Personal or educational use only.  All commercial and enterprise use must be licensed, contact wfz@nimblebooks.com

'''
loop over a folder containing docx and insert them into a single new docx

'''
import os

from docx import Document


def combine_word_documents(folder):
    # Create a new Document where the content will be added
    new_doc = Document()

    # Get all files in the directory
    files = os.listdir(folder)

    # Filter out non-docx files
    docx_files = [f for f in files if f.endswith('.docx')]

    # Sort the files by their creation time
    docx_files.sort(key=lambda x: os.path.getctime(os.path.join(folder, x)))

    # Loop over each .docx file
    for filename in docx_files:
        # Open the .docx file
        doc = Document(os.path.join(folder, filename))

        # Loop over each paragraph in the document
        for paragraph in doc.paragraphs:
            # Add the content of the paragraph into the new document
            new_doc.add_paragraph(paragraph.text)

    # Save the new document in the same directory
    new_doc.save(os.path.join(folder, 'combined4.docx'))


folder = "working/contracted/active_copyedit/book4"
combine_word_documents(folder)

HISTORICAL
RECORD
OF
U.S.GOVERNMENT
INVOLVEMENT
WITH
UNIDENTIFIED
ANOMALOUS
PHENOMENA(UAP)
