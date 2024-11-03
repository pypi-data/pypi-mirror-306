import logging
import re
import traceback
from pathlib import Path

import pandas as pd

from Codexes2Gemini.classes.Codexes.Distributors.LSI import create_LSI_ACS_spreadsheet
from Codexes2Gemini.classes.Codexes.Metadata.Metadatas import Metadatas

# import docx reference enumerations

# ogging.basicConfig(level=print, format='%(asctime)s - %(levelname)s - %(message)s', file='metadatas2output_schemas.log')


try:
    from pypandoc import convert_file
except ImportError as e:
    logging.error("Pandoc not installed. No pypandoc module")
    logging.error(e)
    import pypandoc
    from pypandoc import convert_file
import streamlit as st
# import Inches
from docx.shared import Inches


def read_json_metadatas_file(filename):
    """
    Reads and parses the contents of a JSON metadata file.

    Parameters:
    - filename (str): The name of the file to read.

    Returns:
    - data (dict): The parsed data from the JSON file.
    """
    with open(filename, 'r') as file:
        data = json.load(file)
    return data


def docx_production_df2docx(document, docx_production_df, filename):
    # rowsdoc = Document(filename)
    for index, row in docx_production_df.iterrows():
        logging.debug(f'inserting row {row}')
        if row["style"] == "Section Break":
            try:
                row["string"] = ""
                document.add_section_break()
            except Exception as e:
                errmessage = f"Error adding section break: {e}"
                logging.error(errmessage)

        if row["style"] == "Page Break":
            try:
                document.add_page_break()
                logging.debug('added page break')
            except Exception as e:
                errmessage = str(e) + 'error in adding page break'
                logging.error(errmessage)
            continue
        if row["string"] == "insert ToC here":
            try:
                row["string"] = ""
                document.add_paragraph(row["string"], row["style"])
                logging.debug('added placeholder for toc')
            except Exception as e:
                errmessage = str(e) + 'error in adding toc placeholder'
                logging.error(errmessage)
            continue
        logging.debug("style is", row["style"])
        logging.debug(f"string is {row['string']}")
        # remove newline if it occurs at beginning of string
        row["string"] = re.sub(r'^\n', '', row["string"])
        # if a newline is surrouonded only by spaces,
        # remove the spaces
        row["string"] = re.sub(r'\n\s*\n', '\n\n', row["string"])
        # if two or more newlines are found in a row, convert them to only one
        row["string"] = re.sub(r'\n{2,}', '\n', row["string"])
        # if a newline is found at the end of a string, remove it
        row["string"] = re.sub(r'\n$', '', row["string"])
        logging.debug("Cleaned up newlines")
        if row["style"] == "Image":
            if row["string"] != "Image Placeholder":
                try:
                    document.add_picture(row["string"], width=Inches(5.5))
                    logging.debug('added image')
                except Exception as e:
                    errmessage = str(e) + 'error in adding image'
                    logging.error(errmessage)
                    document.add_paragraph("Image not found, check log to see if it was generated.", row["style"])
            else:
                try:
                    document.add_paragraph("", row["style"])
                except Exception as e:
                    errmessage = str(e) + 'error in adding image placeholder'
                    logging.error(errmessage)
        else:
            # count newlines in the text-bearing string
            newline_count = row["string"].count('\n')
            # now if there are multiple paragraphs in the string, split them and add them as separate paragraphs
            if newline_count >= 1:
                paragraphs = row["string"].split('\n')
                for p in paragraphs:

                    try:
                        document.add_paragraph(p, row["style"])

                    except Exception as e:
                        st.error(e)
                        st.error(row['string'])
                        logging.error(e)
                        continue
                #   print('added paragraphs')
            else:
                try:
                    document.add_paragraph(row['string'], row["style"])
                except Exception as e:
                    st.error(e)
                    st.error(row['string'])
                    logging.error(e)

        if row["string"] == "toc_insert_mark":
            try:
                add_toc(document)
                logging.debug('added toc')
            except Exception as e:
                errmessage = str(e) + 'error in adding toc'
                logging.error(errmessage)
            for i in range(1, 4):
                document.add_heading(f'Heading Level {i}', level=i)
                document.add_paragraph('Some content here...')
        # rint('finished for loop')
        if row["string"] == "":
            continue
            # document.save(filename)
    return document


def merge_documents(doc1, doc2, thisdoc_dir):
    # Load the documents

    # Append content of document1 to document2
    for element in doc1.element.body:
        doc2.element.body.append(element)

    # Save the merged document
    doc2.save(thisdoc_dir + '/merged.docx')

    return doc2


def metadatas2docx(df):
    pass


def metadatas2bookjson(metadatas, thisdoc_dir, distributor="LSI"):
    # convert the metadata to a json file
    st.info('entering metadatas2bookjson')


    bookjsonkeys = ["BookID", "BookTitle", "SubTitle", "Byline", "ImprintText", "ImageFileName", "settings",
                    "distributor", "InvertedColor", "DominantColor", "BaseFont", "trimsizewidth", "trimsizeheight",
                    "spinewidth", "backtext"]
    bookjson = {key: "" for key in bookjsonkeys}
    bookjson["BookID"] = metadatas.get_attribute("ISBN")
    bookjson["BookTitle"] = metadatas.get_attribute("gemini_title")
    bookjson["SubTitle"] = metadatas.get_attribute("subtitle")
    bookjson["Byline"] = metadatas.get_attribute("gemini_authors")
    bookjson["ImprintText"] = metadatas.get_attribute("imprint") + ": " + metadatas.get_attribute("imprint_slug")
    bookjson["ImageFileName"] = ""  # metadatas.get_attribute("cover_image")
    bookjson["settings"] = "default"
    bookjson["distributor"] = "LSI"
    bookjson["InvertedColor"] = "White"
    bookjson["DominantColor"] = "Nimble Maroon"
    bookjson["BaseFont"] = "Skolar PE Regular"
    bookjson["trimsizewidth"] = 8.5  # metadatas.get_attribute("trim_size")
    bookjson["trimsizeheight"] = 11  # metadatas.get_attribute("trim_size")
    bookjson["spinewidth"] = metadatas.get_attribute("spinewidth")
    if metadatas.get_attribute("motivation") is None:
        metadatas.set_attribute("motivation", "")

    backtext = metadatas.get_attribute("gemini_motivation") + '\n' + metadatas.get_attribute("gemini_summary") + '\n' + \
               metadatas.get_attribute('description_of_annotations') + '\n' + metadatas.get_attribute('source_text')
    bookjson["backtext"] = backtext
    # st.write(bookjson)
    logging.debug(str(bookjson))

    try:
        with open(thisdoc_dir + '/' + Path(thisdoc_dir).stem + "_book.json", 'w') as f:

            f.write(json.dump(bookjson, f))


    except Exception as e:
        logging.error(f"Error in saving bookjson: {str(e)}")
        st.error(f"Error in saving bookjson: {str(e)}")

    return bookjson


def metadatas2internationaledition(metadatas, thisdoc_dir, languages, international_presets, edition_name):
    # convert

    # add bespoke content that motivatess the international edition

    # translate selected metadata to other languages
    international_text = 'placeholder'  # hgrtext2googlecloudtranslate.py(metadatas, languages, international_presets)
    # assemble supplementary front matter section

    return "successfully built book json file for Scribus"


import json


def create_section(metadatas, key=None):
    predefined_sections_file_path = metadatas.get("predefined_sections_file_path")
    # print(f"Attempting to open file: {predefined_sections_file_path}")
    metadatas_dict = metadatas.to_dict()
    # print(metadatas_dict.keys())
    try:
        with open(predefined_sections_file_path, 'r') as f:
            file_contents = f.read()
            # print("File contents:")
            # print(file_contents)
            PREDEFINED_SECTIONS = json.loads(file_contents)
            if isinstance(PREDEFINED_SECTIONS, list):
                PREDEFINED_SECTIONS = PREDEFINED_SECTIONS[0]
    except FileNotFoundError:
        print(f"File not found: {predefined_sections_file_path}")
        logging.error(f"File not found: {predefined_sections_file_path}")
        return None
    except json.JSONDecodeError as e:
        logmsg = f"JSON decode error: {e}"
        logmsg += f"Error occurred near: {file_contents[max(0, e.pos - 20):e.pos + 20]}"
        logging.error(logmsg)
        return None
    except Exception as e:
        logging.error(f"Unexpected error: {traceback.format_exc()}")
        return None

    print(PREDEFINED_SECTIONS.keys())
    # print(metadatas.list_all_attributes())
    if key and key in PREDEFINED_SECTIONS:
        print("***" * 5)
        print(f"Section {key} exists in the predefined sections")
        section = PREDEFINED_SECTIONS[key]

        # Format pre_content
        section['pre_content'] = [content.format_map(metadatas_dict) for content in section['pre_content']]
        #  print(f"precontent is {section['pre_content']}")

        # Handle gemini_result
        results = []
        prompt_keys = section['gemini_result']
        for p in prompt_keys:
            print(f"metadata_key is {p}")
            result_value = metadatas.get(p)
            # print(f"result_value for section {key} prompt {p} is {result_value}")
            results.append(result_value)

        section["gemini_result"] = results

        # Format post_content
        section['post_content'] = [content.format_map(metadatas_dict) for content in section['post_content']]

    else:
        print(f"Creating new section {key}")
        section = {
            "title": "",
            "pre_content": [],
            "gemini_result": [],
            "post_content": []
        }
    # print("***" * 10, "SECTION CONTENT")
    # print(section["pre_content"] + section["gemini_result"] + section["post_content"])

    if "gemini_result" in section and section["gemini_result"] is not None and section["gemini_result"] != []:
        logging.info(f'section["gemini_result"] is not empty')
        content = section["pre_content"] + section["gemini_result"] + section["post_content"]
    else:
        logging.error(f'section["gemini_result"] is None or empty')
        content = section["pre_content"] + section["post_content"]

    return {
        "title": section["title"].format_map(metadatas_dict),
        "content": content
    }


def create_sections(metadatas, section_keys):
    sections = []
    print(f"section_keys are {section_keys}")
    for key in section_keys:
        test = create_section(metadatas, key=key)
        with open(metadatas.get("thisdoc_dir") + "/" + key + ".txt", 'w') as f:
            f.write(json.dumps(test))
        sections.append(create_section(metadatas, key=key))
    return sections


def metadatas2markdown(metadatas, thisdoc_dir):
    metadatas.set_attribute("year", "2024")
    # LaTeX declarations
    latex_declarations = """---
title: "{}"
author: {}
header-includes:
  - \\usepackage{{fancyhdr}}
  - \\pagestyle{{fancy}}
  - \\fancyhf{{}}
  - \\fancyfoot[C]{{\\thepage}}
  - \\usepackage{{longtable}}

output: pdf_document
---
\\pagenumbering{{roman}}
\\newpage
""".format(metadatas.get("gemini_title"), metadatas.get("gemini_authors"))

    if "page-by-page" in metadatas.get("flags"):
        df = pd.read_csv(thisdoc_dir + '/page_by_page_results_df.csv')
        print(metadatas["page_by_page_presets"])
        latex_tables = generate_latex_tables_from_df(df, metadatas["page_by_page_presets"])
    else:
        logging.info("pagebypage is false, not creating Latex tables")
        latex_tables = []
    print(metadatas.get("section_keys"))
    sections = create_sections(metadatas, metadatas.get("section_keys"))
    print(sections)
    if "page-by-page" in metadatas.get('flags'):
        if latex_tables:
            sections.extend(latex_tables)
        else:
            logging.error("Latex tables are missing, not adding to markdown")

    # Convert sections to markdown
    markdown_content = latex_declarations
    for section in sections:
        if "title" in section:
            if section["title"]:
                markdown_content += f"\n\n# {section['title']}\n\n"
            for line in section["content"]:
                if line:
                    markdown_content += f"{line}\n\n"
            markdown_content += "\\newpage"
        else:
            logging.error(f"no title element for section {section}, skipping")

    # Write to file
    output_file_path = f"{thisdoc_dir}/output.md"
    with open(output_file_path, "w") as f:
        f.write(markdown_content)

    logging.info(f'Saved Pandoc Markdown file to {output_file_path}')
    st.info(f'Saved Pandoc Markdown file to {output_file_path}')

    return markdown_content

    i


def add_toc(document, title='Annotations'):
    from docx.oxml import OxmlElement
    # Add a title for the ToC
    try:
        document.add_paragraph(title, style='TOC Heading')

        # Create the ToC field code element
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
        run._r.append(fldSimple)
        logging.info("inserted ToC")
    except Exception as e:
        logging.error(f"Error in adding toc: {str(e)}")
    return


def generate_markdown_tables_from_df(filepath, pbp_presets):
    """
    Generates markdown tables from a DataFrame with columns in pairs.

    Parameters:
    df (pd.DataFrame): The DataFrame containing the data.

    Returns:
    str: A string containing all the markdown tables.
    """
    df = pd.read_csv(filepath)

    # drop rows that have Nan values
    df = df.dropna()

    # Function to generate a markdown table for a pair of columns
    def generate_markdown_table(df, col1, col2, title):
        markdown = f"### {title}\n\n"
        markdown += "| Page Number | Text |\n"
        markdown += "| --- | --- |\n"
        for index, row in df.iterrows():
            markdown += f"| {row[col1]} | {row[col2]} |\n"
        return markdown


def generate_latex_tables_from_df(df, pbp_presets):
    """
    Generates LaTeX tables from a DataFrame with columns in pairs.
    Parameters:
    df (pd.DataFrame): The DataFrame containing the data.
    Returns:
    str: A string containing all the LaTeX tables.
    """
    df.dropna(inplace=True)

    # Function to generate a LaTeX table for a pair of columns
    def generate_latex_table(df, col1, col2, title):
        # latex = f"\\section*{{{title}}}\n"
        latex = "\\begin{longtable}{|c|p{10cm}|}\n\\hline\n"
        latex += "Page Number & Text \\\\\n\\hline\n"
        latex += "\\endhead\n"
        for index, row in df.iterrows():
            latex += f"{row[col1]} & {row[col2]} \\\\\n\\hline\n"
        latex += "\\end{longtable}\n"
        return latex

    # List to store LaTeX tables
    latex_tables = []
    # Loop through the columns in pairs
    for i in range(0, len(df.columns), 2):
        col1 = df.columns[i]
        col2 = df.columns[i + 1]
        # if i is even
        if i % 2 == 0:
            title = pbp_presets[i // 2]
            latex = generate_latex_table(df, col1, col2, title)
            latex_table_dict = {"title": title, "content": [latex]}
            latex_tables.append(latex_table_dict)
    return latex_tables


if __name__ == "__main__":
    md = Metadatas()
    metadatas_file = "test/json/edited_df.json"
    metadatas = read_json_metadatas_file(metadatas_file)
    # logging.debug(type(metadatas))
    # hoist inner dictionary to next level up
    metadatas = metadatas["0"]
    thisdoc_dir = metadatas["thisdoc_dir"]

    bookjson = metadatas2bookjson(metadatas, thisdoc_dir)
    lsi_df = create_LSI_ACS_spreadsheet(metadatas)
    print(lsi_df.head())
    print('head')
    try:
        lsi_df.to_csv('/Users/fred/test.csv', index=False)
        print(f'saved ACS spreadsheet to output')
    except Exception as e:
        print(f"error {e} saving LSI spreadsheet")
